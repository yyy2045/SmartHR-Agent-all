from __future__ import annotations

import json
import uuid
from collections import Counter
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Literal

from fastapi import UploadFile
from pydantic import BaseModel, ConfigDict, Field, model_validator
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool
from starlette.datastructures import Headers

from app.database import Base
from app.models import (
    ApplicationResumeDocument,
    Candidate,
    HardRequirement,
    Job,
    JobApplication,
    JobCriteriaVersion,
    ResumeDocument,
    ResumeRedaction,
    ResumeTextSegment,
    ScoringDimension,
    ScreeningBatch,
    ScreeningResult,
    User,
)
from app.schemas.screening import ResumeAnalysisDraft
from app.services.file_storage import resolve_private_file, store_resume_upload
from app.services.model_payload import build_resume_model_payload
from app.services.resume_analysis import analyze_resume_document, screening_result_sort_key
from app.services.resume_parser import OCRLine, ParseResult, parse_resume_file
from app.services.resume_redactor import redact_resume_segments
from app.services.security import hash_password

ResumeFormat = Literal["pdf", "docx", "scanned_pdf", "jpg", "png"]
Language = Literal["zh-CN", "en-US"]
HardStatus = Literal["passed", "failed", "unknown"]
ExpectedGroup = Literal["passed", "low_match", "auto_rejected"]
Scenario = Literal[
    "high_match",
    "low_match",
    "hard_failure",
    "missing_information",
    "ambiguous_context",
]

EXPECTED_REDACTION_TYPES = {
    "name",
    "phone",
    "email",
    "id_number",
    "address",
    "social_account",
}
FORMAT_EXTENSIONS: dict[ResumeFormat, str] = {
    "pdf": ".pdf",
    "docx": ".docx",
    "scanned_pdf": ".pdf",
    "jpg": ".jpg",
    "png": ".png",
}
FORMAT_MIME_TYPES: dict[ResumeFormat, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "scanned_pdf": "application/pdf",
    "jpg": "image/jpeg",
    "png": "image/png",
}


class DimensionSpec(BaseModel):
    model_config = ConfigDict(extra="forbid")

    name: str
    description: str
    weight_percent: int = Field(ge=0, le=100)


class JobSpec(BaseModel):
    model_config = ConfigDict(extra="forbid")

    key: str
    title: str
    department: str
    hard_requirement_title: str
    hard_requirement_value: str
    dimensions: list[DimensionSpec]
    evidence_zh: str
    evidence_en: str

    @model_validator(mode="after")
    def validate_weights(self) -> JobSpec:
        if sum(item.weight_percent for item in self.dimensions) != 100:
            raise ValueError("职位评分维度权重必须合计 100")
        return self


class ResumeSpec(BaseModel):
    model_config = ConfigDict(extra="forbid")

    id: str
    job_key: str
    language: Language
    format: ResumeFormat
    scenario: Scenario
    hard_status: HardStatus
    score: int = Field(ge=0, le=100)
    expected_group: ExpectedGroup


class EvaluationDataset(BaseModel):
    model_config = ConfigDict(extra="forbid")

    version: str
    jobs: list[JobSpec]
    resumes: list[ResumeSpec]

    @model_validator(mode="after")
    def validate_references(self) -> EvaluationDataset:
        job_keys = [job.key for job in self.jobs]
        if len(job_keys) != len(set(job_keys)):
            raise ValueError("评测职位 key 不能重复")
        resume_ids = [resume.id for resume in self.resumes]
        if len(resume_ids) != len(set(resume_ids)):
            raise ValueError("评测简历 ID 不能重复")
        unknown_jobs = {resume.job_key for resume in self.resumes} - set(job_keys)
        if unknown_jobs:
            raise ValueError(f"评测简历引用了未知职位：{sorted(unknown_jobs)}")
        return self


class EvaluationReport(BaseModel):
    dataset_version: str
    passed: bool
    job_count: int
    resume_count: int
    language_counts: dict[str, int]
    format_counts: dict[str, int]
    scenario_counts: dict[str, int]
    expected_group_counts: dict[str, int]
    actual_group_counts: dict[str, int]
    redaction_types: list[str]
    generated_file_count: int
    completed_analysis_count: int
    payload_leak_count: int
    issues: list[str]


@dataclass(frozen=True)
class SyntheticIdentity:
    name: str
    phone: str
    email: str
    id_number: str
    address: str
    social_account: str

    @property
    def sensitive_values(self) -> tuple[str, ...]:
        return (
            self.name,
            self.phone,
            self.email,
            self.id_number,
            self.address,
            self.social_account,
        )


@dataclass(frozen=True)
class EvaluationSegment:
    segment_key: str
    normalized_text: str


@dataclass(frozen=True)
class PreparedResume:
    spec: ResumeSpec
    identity: SyntheticIdentity
    document_id: uuid.UUID
    parse_result: ParseResult
    expected_evidence: str
    hard_evidence: str | None
    redaction_types: frozenset[str]


@dataclass(frozen=True)
class JobRuntime:
    criteria_id: uuid.UUID
    requirement_id: uuid.UUID
    dimension_ids: tuple[uuid.UUID, ...]


class StaticOCREngine:
    def __init__(self, text: str) -> None:
        self.text = text

    def recognize(self, _: bytes) -> list[OCRLine]:
        return [
            OCRLine(text=line, confidence=0.99)
            for line in self.text.splitlines()
            if line.strip()
        ]


class DeterministicEvaluationClient:
    model = "mvp-evaluation-stub"

    def __init__(
        self,
        *,
        samples_by_candidate: dict[str, PreparedResume],
        jobs: dict[str, JobSpec],
        runtimes: dict[str, JobRuntime],
    ) -> None:
        self.samples_by_candidate = samples_by_candidate
        self.jobs = jobs
        self.runtimes = runtimes
        self.payloads: list[dict[str, object]] = []

    async def analyze_resume(
        self,
        payload: dict[str, object],
        *,
        validation_feedback: str | None = None,
        previous_analysis: dict[str, object] | None = None,
    ) -> ResumeAnalysisDraft:
        del validation_feedback, previous_analysis
        self.payloads.append(payload)
        candidate_code = str(payload["candidate_code"])
        prepared = self.samples_by_candidate[candidate_code]
        spec = prepared.spec
        job = self.jobs[spec.job_key]
        runtime = self.runtimes[spec.job_key]
        evidence = _evidence_reference(payload, prepared.expected_evidence)
        hard_evidence = (
            [_evidence_reference(payload, prepared.hard_evidence)]
            if prepared.hard_evidence is not None
            else []
        )
        hard_rationale = {
            "passed": "合成样本明确满足最低工作年限。",
            "failed": "合成样本明确低于最低工作年限。",
            "unknown": "合成样本未提供可确认的工作年限。",
        }[spec.hard_status]
        return ResumeAnalysisDraft.model_validate(
            {
                "candidate_profile": {
                    "education": [],
                    "work_experiences": [],
                    "projects": [],
                    "skills": [
                        {
                            "name": job.dimensions[0].name,
                            "level": "合成评测",
                            "evidence": [evidence],
                        }
                    ],
                    "certifications": [],
                    "languages": [],
                },
                "hard_requirements": [
                    {
                        "requirement_id": str(runtime.requirement_id),
                        "status": spec.hard_status,
                        "rationale": hard_rationale,
                        "evidence": hard_evidence,
                    }
                ],
                "dimension_scores": [
                    {
                        "dimension_id": str(dimension_id),
                        "score": spec.score,
                        "rationale": f"固定合成样本评分：{spec.score}",
                        "missing_items": [],
                        "evidence": [evidence],
                    }
                    for dimension_id in runtime.dimension_ids
                ],
                "strengths": [prepared.expected_evidence]
                if spec.score >= 70
                else [],
                "gaps": ["综合匹配度偏低"] if spec.score < 70 else [],
                "missing_items": ["工作年限"]
                if spec.hard_status == "unknown"
                else [],
                "interview_questions": ["请进一步说明相关项目中的具体职责。"],
            }
        )


def load_dataset() -> EvaluationDataset:
    path = Path(__file__).with_name("dataset.json")
    return EvaluationDataset.model_validate_json(path.read_text(encoding="utf-8"))


def _synthetic_identity(index: int, language: Language) -> SyntheticIdentity:
    if language == "zh-CN":
        chinese_names = (
            "星河甲",
            "星河乙",
            "星河丙",
            "星河丁",
            "星河戊",
            "星河己",
            "星河庚",
            "星河辛",
            "星河壬",
            "星河癸",
            "云岚甲",
            "云岚乙",
            "云岚丙",
            "云岚丁",
            "云岚戊",
        )
        return SyntheticIdentity(
            name=chinese_names[(index - 1) // 2],
            phone=f"1300000{index:04d}",
            email=f"synthetic{index:02d}@example.test",
            id_number=f"99010119900101{index:04d}",
            address=f"北京市海淀区合成路{100 + index}号",
            social_account=f"synthetic-cn-{index:02d}",
        )
    english_names = (
        "Avery Synthetic",
        "Blake Synthetic",
        "Casey Synthetic",
        "Devon Synthetic",
        "Emery Synthetic",
        "Finley Synthetic",
        "Gray Synthetic",
        "Harper Synthetic",
        "Indigo Synthetic",
        "Jordan Synthetic",
        "Kai Synthetic",
        "Logan Synthetic",
        "Morgan Synthetic",
        "Nova Synthetic",
        "Parker Synthetic",
    )
    return SyntheticIdentity(
        name=english_names[(index - 2) // 2],
        phone=f"+1 202-555-{1000 + index}",
        email=f"synthetic{index:02d}@example.test",
        id_number=f"000-00-{1000 + index}",
        address=f"{100 + index} Synthetic Road, Apt {index}",
        social_account=f"synthetic-en-{index:02d}",
    )


def _hard_evidence(spec: ResumeSpec) -> str | None:
    if spec.hard_status == "unknown":
        return None
    if spec.language == "zh-CN":
        return "5 年相关经验" if spec.hard_status == "passed" else "1 年相关经验"
    return (
        "5 years of relevant experience"
        if spec.hard_status == "passed"
        else "1 year of relevant experience"
    )


def _resume_text(
    spec: ResumeSpec,
    job: JobSpec,
    identity: SyntheticIdentity,
) -> tuple[str, str, str | None]:
    evidence = job.evidence_zh if spec.language == "zh-CN" else job.evidence_en
    hard_evidence = _hard_evidence(spec)
    if spec.language == "zh-CN":
        experience = (
            f"相关经验：{hard_evidence}"
            if hard_evidence is not None
            else "相关经验：年限信息未提供"
        )
        text = "\n".join(
            (
                f"姓名：{identity.name}",
                f"电话：{identity.phone}",
                f"邮箱：{identity.email}",
                f"身份证号：{identity.id_number}",
                f"现居地：{identity.address}",
                f"Github：{identity.social_account}",
                experience,
                f"核心经历：{evidence}",
                f"评测场景：{spec.scenario}",
                "本简历为 SmartHR MVP 自动生成的完全合成样本，不对应任何真实个人。",
            )
        )
    else:
        experience = (
            f"Experience: {hard_evidence}"
            if hard_evidence is not None
            else "Experience: duration not stated"
        )
        text = "\n".join(
            (
                f"Name: {identity.name}",
                f"Phone: {identity.phone}",
                f"Email: {identity.email}",
                f"SSN: {identity.id_number}",
                f"Address: {identity.address}",
                f"LinkedIn: {identity.social_account}",
                experience,
                f"Core evidence: {evidence}",
                f"Evaluation scenario: {spec.scenario}",
                "This resume is a fully synthetic SmartHR MVP sample "
                "and represents no real person.",
            )
        )
    return text, evidence, hard_evidence


def _insert_text(page: object, text: str) -> None:
    import fitz

    contains_chinese = any("\u4e00" <= character <= "\u9fff" for character in text)
    font_name = "china-s" if contains_chinese else "helv"
    page.insert_textbox(
        fitz.Rect(36, 36, 559, 806),
        text,
        fontsize=11,
        fontname=font_name,
        lineheight=1.35,
    )


def _text_image_bytes(text: str) -> bytes:
    import fitz

    document = fitz.open()
    try:
        page = document.new_page(width=595, height=842)
        _insert_text(page, text)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        return pixmap.tobytes("png")
    finally:
        document.close()


def _write_sample_file(path: Path, sample_format: ResumeFormat, text: str) -> None:
    if sample_format == "docx":
        from docx import Document

        document = Document()
        for line in text.splitlines():
            document.add_paragraph(line)
        document.save(path)
        return

    if sample_format == "pdf":
        import fitz

        document = fitz.open()
        try:
            page = document.new_page(width=595, height=842)
            _insert_text(page, text)
            document.save(path)
        finally:
            document.close()
        return

    image_bytes = _text_image_bytes(text)
    if sample_format == "scanned_pdf":
        import fitz

        document = fitz.open()
        try:
            page = document.new_page(width=595, height=842)
            page.insert_image(page.rect, stream=image_bytes)
            document.save(path)
        finally:
            document.close()
        return

    from PIL import Image

    with Image.open(BytesIO(image_bytes)) as image:
        if sample_format == "jpg":
            image.convert("RGB").save(path, format="JPEG", quality=92)
        else:
            image.save(path, format="PNG")


def _evidence_reference(payload: dict[str, object], quote: str | None) -> dict[str, str]:
    if quote is None:
        raise ValueError("明确判断缺少评测证据")
    segments = payload.get("segments")
    if not isinstance(segments, list):
        raise ValueError("模型载荷缺少片段")
    for segment in segments:
        if not isinstance(segment, dict):
            continue
        text = str(segment.get("text", ""))
        if quote in text:
            return {"segment_key": str(segment["segment_key"]), "quote": quote}
    raise ValueError(f"评测证据不在模型载荷中：{quote}")


async def _prepare_resume(
    *,
    index: int,
    spec: ResumeSpec,
    job: JobSpec,
    source_root: Path,
    storage_root: Path,
) -> tuple[PreparedResume, list[ResumeTextSegment]]:
    identity = _synthetic_identity(index, spec.language)
    text, evidence, hard_evidence = _resume_text(spec, job, identity)
    extension = FORMAT_EXTENSIONS[spec.format]
    source_path = source_root / f"{spec.id.lower()}{extension}"
    _write_sample_file(source_path, spec.format, text)

    with source_path.open("rb") as source:
        upload = UploadFile(
            filename=source_path.name,
            file=source,
            headers=Headers({"content-type": FORMAT_MIME_TYPES[spec.format]}),
        )
        stored = await store_resume_upload(
            upload,
            storage_root=storage_root,
            job_id=uuid.uuid5(uuid.NAMESPACE_URL, f"job:{spec.job_key}"),
            batch_id=uuid.uuid5(uuid.NAMESPACE_URL, f"batch:{spec.job_key}"),
            max_size_bytes=20 * 1024 * 1024,
        )
    stored_path = resolve_private_file(storage_root, stored.storage_key)
    ocr_engine = (
        StaticOCREngine(text)
        if spec.format in {"scanned_pdf", "jpg", "png"}
        else None
    )
    parse_result = parse_resume_file(
        stored_path,
        stored.detected_type,
        ocr_engine=ocr_engine,
    )
    expected_method = {
        "pdf": "pdf_text",
        "docx": "docx_text",
        "scanned_pdf": "pdf_ocr",
        "jpg": "image_ocr",
        "png": "image_ocr",
    }[spec.format]
    if parse_result.extraction_method != expected_method:
        raise ValueError(
            f"{spec.id} 解析方式错误：{parse_result.extraction_method} != {expected_method}"
        )

    evaluation_segments = [
        EvaluationSegment(
            segment_key=f"SEG-{position:04d}",
            normalized_text=segment.normalized_text,
        )
        for position, segment in enumerate(parse_result.segments, start=1)
    ]
    document_id = uuid.uuid5(uuid.NAMESPACE_URL, f"resume:{spec.id}")
    redaction = redact_resume_segments(
        f"CAND-{document_id.hex[:12].upper()}",
        evaluation_segments,
    )
    redacted_by_key = {segment.segment_key: segment for segment in redaction.segments}
    redaction_types = {
        match.entity_type
        for segment in redaction.segments
        for match in segment.matches
    }
    rows: list[ResumeTextSegment] = []
    for position, parsed in enumerate(parse_result.segments, start=1):
        segment_key = f"SEG-{position:04d}"
        redacted = redacted_by_key[segment_key]
        row = ResumeTextSegment(
            segment_key=segment_key,
            source_type=parsed.source_type,
            source_index=parsed.source_index,
            page_number=parsed.page_number,
            paragraph_index=parsed.paragraph_index,
            raw_text=parsed.raw_text,
            normalized_text=parsed.normalized_text,
            redacted_text=redacted.redacted_text,
            ocr_confidence=parsed.ocr_confidence,
            sort_order=position - 1,
        )
        row.redactions = [
            ResumeRedaction(
                entity_type=match.entity_type,
                original_text=match.original_text,
                replacement_text=match.replacement_text,
                start_offset=match.start_offset,
                end_offset=match.end_offset,
            )
            for match in redacted.matches
        ]
        rows.append(row)
    return (
        PreparedResume(
            spec=spec,
            identity=identity,
            document_id=document_id,
            parse_result=parse_result,
            expected_evidence=evidence,
            hard_evidence=hard_evidence,
            redaction_types=frozenset(redaction_types),
        ),
        rows,
    )


def _create_evaluation_database() -> tuple[sessionmaker[Session], object]:
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    return sessionmaker(bind=engine, autoflush=False, expire_on_commit=False), engine


def _seed_jobs(
    session_factory: sessionmaker[Session],
    dataset: EvaluationDataset,
) -> dict[str, JobRuntime]:
    runtimes: dict[str, JobRuntime] = {}
    with session_factory() as db:
        user = User(
            username="mvp-evaluator",
            password_hash=hash_password("synthetic-evaluation-only"),
            display_name="MVP 合成评测",
        )
        db.add(user)
        db.flush()
        for job_spec in dataset.jobs:
            job = Job(
                id=uuid.uuid5(uuid.NAMESPACE_URL, f"evaluation-job:{job_spec.key}"),
                owner_id=user.id,
                title=job_spec.title,
                department=job_spec.department,
                original_jd=f"{job_spec.title} MVP 固定合成评测职位",
            )
            criteria = JobCriteriaVersion(
                id=uuid.uuid5(uuid.NAMESPACE_URL, f"evaluation-criteria:{job_spec.key}"),
                job=job,
                version_number=1,
                status="confirmed",
                pass_threshold=70,
                confirmed_by_id=user.id,
                confirmed_at=datetime.now(UTC),
            )
            requirement = HardRequirement(
                id=uuid.uuid5(uuid.NAMESPACE_URL, f"evaluation-requirement:{job_spec.key}"),
                requirement_type="min_experience_years",
                title=job_spec.hard_requirement_title,
                expected_value=job_spec.hard_requirement_value,
                auto_reject=True,
                sort_order=0,
            )
            criteria.hard_requirements = [requirement]
            criteria.scoring_dimensions = [
                ScoringDimension(
                    id=uuid.uuid5(
                        uuid.NAMESPACE_URL,
                        f"evaluation-dimension:{job_spec.key}:{position}",
                    ),
                    name=dimension.name,
                    description=dimension.description,
                    weight_percent=dimension.weight_percent,
                    sort_order=position,
                )
                for position, dimension in enumerate(job_spec.dimensions)
            ]
            batch = ScreeningBatch(
                id=uuid.uuid5(uuid.NAMESPACE_URL, f"evaluation-batch:{job_spec.key}"),
                job=job,
                criteria_version=criteria,
                name=f"{job_spec.title} 固定评测批次",
                status="completed",
                ai_input_mode="redacted",
            )
            db.add(batch)
            db.flush()
            runtimes[job_spec.key] = JobRuntime(
                criteria_id=criteria.id,
                requirement_id=requirement.id,
                dimension_ids=tuple(item.id for item in criteria.scoring_dimensions),
            )
        db.commit()
    return runtimes


def _persist_prepared_resume(
    session_factory: sessionmaker[Session],
    prepared: PreparedResume,
    segment_rows: list[ResumeTextSegment],
) -> str:
    with session_factory() as db:
        batch_id = uuid.uuid5(
            uuid.NAMESPACE_URL,
            f"evaluation-batch:{prepared.spec.job_key}",
        )
        candidate = Candidate(
            id=uuid.uuid5(uuid.NAMESPACE_URL, f"evaluation-candidate:{prepared.spec.id}")
        )
        application = JobApplication(
            id=uuid.uuid5(uuid.NAMESPACE_URL, f"evaluation-application:{prepared.spec.id}"),
            candidate=candidate,
            job_id=uuid.uuid5(
                uuid.NAMESPACE_URL,
                f"evaluation-job:{prepared.spec.job_key}",
            ),
        )
        document = ResumeDocument(
            id=prepared.document_id,
            batch_id=batch_id,
            candidate=candidate,
            application=application,
            original_filename=f"{prepared.spec.id}{FORMAT_EXTENSIONS[prepared.spec.format]}",
            file_extension=FORMAT_EXTENSIONS[prepared.spec.format],
            content_type=FORMAT_MIME_TYPES[prepared.spec.format],
            detected_type="pdf"
            if prepared.spec.format in {"pdf", "scanned_pdf"}
            else prepared.spec.format,
            size_bytes=1,
            storage_key=f"synthetic/{prepared.spec.id.lower()}{FORMAT_EXTENSIONS[prepared.spec.format]}",
            extraction_method=prepared.parse_result.extraction_method,
            segment_count=len(segment_rows),
            text_character_count=sum(len(item.normalized_text) for item in segment_rows),
            redaction_count=sum(len(item.redactions) for item in segment_rows),
            status="completed",
            parsed_at=datetime.now(UTC),
            redacted_at=datetime.now(UTC),
        )
        document.text_segments = segment_rows
        db.add(document)
        db.flush()
        db.add(
            ApplicationResumeDocument(
                application_id=application.id,
                document_id=document.id,
            )
        )
        application.primary_document_id = document.id
        db.commit()
        return document.candidate_code


def _dataset_coverage_issues(dataset: EvaluationDataset) -> list[str]:
    issues: list[str] = []
    if len(dataset.jobs) != 3:
        issues.append(f"固定职位数量应为 3，实际为 {len(dataset.jobs)}")
    if len(dataset.resumes) != 30:
        issues.append(f"固定简历数量应为 30，实际为 {len(dataset.resumes)}")
    formats = Counter(item.format for item in dataset.resumes)
    for sample_format in FORMAT_EXTENSIONS:
        if formats[sample_format] == 0:
            issues.append(f"固定评测集缺少格式：{sample_format}")
    languages = Counter(item.language for item in dataset.resumes)
    for language in ("zh-CN", "en-US"):
        if languages[language] == 0:
            issues.append(f"固定评测集缺少语言：{language}")
    scenarios = Counter(item.scenario for item in dataset.resumes)
    for scenario in (
        "high_match",
        "low_match",
        "hard_failure",
        "missing_information",
        "ambiguous_context",
    ):
        if scenarios[scenario] == 0:
            issues.append(f"固定评测集缺少场景：{scenario}")
    return issues


async def run_mvp_evaluation() -> EvaluationReport:
    dataset = load_dataset()
    issues = _dataset_coverage_issues(dataset)
    jobs = {job.key: job for job in dataset.jobs}
    prepared_resumes: list[PreparedResume] = []
    redaction_types: set[str] = set()
    session_factory, engine = _create_evaluation_database()
    runtimes = _seed_jobs(session_factory, dataset)

    try:
        with TemporaryDirectory(prefix="smarthr-mvp-evaluation-") as directory:
            root = Path(directory)
            source_root = root / "source"
            storage_root = root / "storage"
            source_root.mkdir()
            storage_root.mkdir()
            samples_by_candidate: dict[str, PreparedResume] = {}
            for index, spec in enumerate(dataset.resumes, start=1):
                prepared, segment_rows = await _prepare_resume(
                    index=index,
                    spec=spec,
                    job=jobs[spec.job_key],
                    source_root=source_root,
                    storage_root=storage_root,
                )
                prepared_resumes.append(prepared)
                redaction_types.update(prepared.redaction_types)
                missing_redactions = EXPECTED_REDACTION_TYPES - prepared.redaction_types
                if missing_redactions:
                    issues.append(
                        f"{spec.id} 缺少脱敏类型：{sorted(missing_redactions)}"
                    )
                candidate_code = _persist_prepared_resume(
                    session_factory,
                    prepared,
                    segment_rows,
                )
                samples_by_candidate[candidate_code] = prepared

            client = DeterministicEvaluationClient(
                samples_by_candidate=samples_by_candidate,
                jobs=jobs,
                runtimes=runtimes,
            )
            actual_groups: Counter[str] = Counter()
            completed_count = 0
            for prepared in prepared_resumes:
                response = await analyze_resume_document(
                    prepared.document_id,
                    criteria_version_id=runtimes[prepared.spec.job_key].criteria_id,
                    session_factory=session_factory,
                    ai_client=client,
                )
                if response.get("status") != "completed":
                    issues.append(f"{prepared.spec.id} 分析失败：{response}")
                    continue
                completed_count += 1
                actual_group = str(response.get("group"))
                actual_groups[actual_group] += 1
                if actual_group != prepared.spec.expected_group:
                    issues.append(
                        f"{prepared.spec.id} 分组错误：{actual_group} != "
                        f"{prepared.spec.expected_group}"
                    )
                if prepared.spec.hard_status == "unknown" and actual_group == "auto_rejected":
                    issues.append(f"{prepared.spec.id} 信息缺失被错误自动淘汰")
                if (
                    prepared.spec.scenario == "low_match"
                    and prepared.spec.hard_status != "failed"
                    and actual_group == "auto_rejected"
                ):
                    issues.append(f"{prepared.spec.id} 低匹配分被错误自动淘汰")

            payload_leak_count = 0
            serialized_payloads = [
                json.dumps(payload, ensure_ascii=False) for payload in client.payloads
            ]
            for prepared, serialized in zip(
                prepared_resumes,
                serialized_payloads,
                strict=True,
            ):
                leaked_values = [
                    value
                    for value in prepared.identity.sensitive_values
                    if value and value in serialized
                ]
                forbidden_file_values = (
                    f"{prepared.spec.id}{FORMAT_EXTENSIONS[prepared.spec.format]}",
                    f"synthetic/{prepared.spec.id.lower()}",
                )
                leaked_values.extend(
                    value for value in forbidden_file_values if value in serialized
                )
                if leaked_values:
                    payload_leak_count += 1
                    issues.append(
                        f"{prepared.spec.id} 模型载荷包含禁止内容：{leaked_values}"
                    )

            with session_factory() as db:
                for job_key in jobs:
                    batch_id = uuid.uuid5(
                        uuid.NAMESPACE_URL,
                        f"evaluation-batch:{job_key}",
                    )
                    results = list(
                        db.scalars(
                            select(ScreeningResult)
                            .join(ResumeDocument)
                            .where(ResumeDocument.batch_id == batch_id)
                        )
                    )
                    ordered = sorted(results, key=screening_result_sort_key)
                    group_order = {
                        "passed": 0,
                        "low_match": 1,
                        "auto_rejected": 2,
                    }
                    observed = [group_order[item.ai_group or "auto_rejected"] for item in ordered]
                    if observed != sorted(observed):
                        issues.append(f"{job_key} 默认排序没有按业务分组稳定排列")
                for document in db.scalars(select(ResumeDocument)):
                    build_resume_model_payload(document)

            expected_groups = Counter(item.expected_group for item in dataset.resumes)
            return EvaluationReport(
                dataset_version=dataset.version,
                passed=not issues,
                job_count=len(dataset.jobs),
                resume_count=len(dataset.resumes),
                language_counts=dict(Counter(item.language for item in dataset.resumes)),
                format_counts=dict(Counter(item.format for item in dataset.resumes)),
                scenario_counts=dict(Counter(item.scenario for item in dataset.resumes)),
                expected_group_counts=dict(expected_groups),
                actual_group_counts=dict(actual_groups),
                redaction_types=sorted(redaction_types),
                generated_file_count=len(prepared_resumes),
                completed_analysis_count=completed_count,
                payload_leak_count=payload_leak_count,
                issues=issues,
            )
    finally:
        Base.metadata.drop_all(engine)
        engine.dispose()
