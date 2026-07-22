from pathlib import Path

import fitz
import pytest
from docx import Document
from PIL import Image

from app.services.resume_parser import (
    OCRLine,
    ResumeParseError,
    normalize_resume_text,
    parse_resume_file,
)


class FakeOCREngine:
    def __init__(self, lines: list[OCRLine]) -> None:
        self.lines = lines
        self.calls = 0

    def recognize(self, image_bytes: bytes) -> list[OCRLine]:
        assert image_bytes
        self.calls += 1
        return self.lines


def create_pdf(path: Path, text: str = "") -> None:
    document = fitz.open()
    page = document.new_page()
    if text:
        page.insert_text((72, 72), text)
    document.save(path)
    document.close()


def test_normalize_resume_text_removes_control_and_collapses_space() -> None:
    value = "  Senior\tEngineer  \r\n\r\n\x00  Python   FastAPI  "

    assert normalize_resume_text(value) == "Senior Engineer\n\nPython FastAPI"


def test_electronic_pdf_prefers_direct_text(tmp_path: Path) -> None:
    path = tmp_path / "resume.pdf"
    create_pdf(
        path,
        "Senior backend engineer with Python FastAPI PostgreSQL Redis Docker "
        "and distributed systems experience.",
    )
    ocr = FakeOCREngine([OCRLine("should not run")])

    result = parse_resume_file(path, "pdf", ocr_engine=ocr, min_pdf_text_characters=20)

    assert result.extraction_method == "pdf_text"
    assert result.segments[0].page_number == 1
    assert "backend engineer" in result.segments[0].normalized_text
    assert ocr.calls == 0


def test_pdf_with_insufficient_text_falls_back_to_ocr(tmp_path: Path) -> None:
    path = tmp_path / "scan.pdf"
    create_pdf(path, "short")
    ocr = FakeOCREngine(
        [
            OCRLine("张三", 0.98),
            OCRLine("Python 后端工程师", 0.92),
        ]
    )

    result = parse_resume_file(path, "pdf", ocr_engine=ocr, min_pdf_text_characters=80)

    assert result.extraction_method == "pdf_ocr"
    assert result.segments[0].normalized_text == "张三\nPython 后端工程师"
    assert result.segments[0].ocr_confidence == pytest.approx(0.95)
    assert ocr.calls == 1


def test_docx_keeps_paragraph_positions_and_table_text(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("候选人简介")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "FastAPI"
    document.add_paragraph("项目经历")
    document.save(path)

    result = parse_resume_file(path, "docx")

    assert result.extraction_method == "docx_text"
    assert [item.normalized_text for item in result.segments] == [
        "候选人简介",
        "Python",
        "FastAPI",
        "项目经历",
    ]
    assert [item.paragraph_index for item in result.segments] == [1, 2, 3, 4]


def test_image_uses_ocr_and_empty_result_is_reported(tmp_path: Path) -> None:
    path = tmp_path / "resume.png"
    Image.new("RGB", (20, 20), "white").save(path)
    ocr = FakeOCREngine([OCRLine("Data Analyst", 0.88)])

    result = parse_resume_file(path, "png", ocr_engine=ocr)

    assert result.extraction_method == "image_ocr"
    assert result.segments[0].source_type == "image_ocr"
    assert result.segments[0].normalized_text == "Data Analyst"

    with pytest.raises(ResumeParseError, match="未识别到有效文本") as error:
        parse_resume_file(path, "png", ocr_engine=FakeOCREngine([]))
    assert error.value.code == "empty_text"


def test_encrypted_and_damaged_pdf_have_clear_errors(tmp_path: Path) -> None:
    encrypted_path = tmp_path / "encrypted.pdf"
    document = fitz.open()
    document.new_page().insert_text((72, 72), "private resume")
    document.save(
        encrypted_path,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner-password",
        user_pw="user-password",
    )
    document.close()

    with pytest.raises(ResumeParseError, match="已加密") as encrypted_error:
        parse_resume_file(encrypted_path, "pdf", ocr_engine=FakeOCREngine([]))
    assert encrypted_error.value.code == "encrypted_pdf"

    damaged_path = tmp_path / "damaged.pdf"
    damaged_path.write_bytes(b"%PDF-damaged")
    with pytest.raises(ResumeParseError, match="损坏") as damaged_error:
        parse_resume_file(damaged_path, "pdf", ocr_engine=FakeOCREngine([]))
    assert damaged_error.value.code == "invalid_pdf"
